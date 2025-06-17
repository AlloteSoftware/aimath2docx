import html.entities
import re
import sys
import latex2mathml.converter
import mathml2omml # type: ignore
from docx import Document
from docx.oxml import parse_xml
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
import regex
import csv
from io import StringIO
import xml.etree.ElementTree as ET

def convert_file(input_path, output_path):
    try:
        with open(input_path, "r", encoding="utf-8") as f:
            markdown = f.read()
        result = markdown_to_docx(markdown, output_path)
        return bool(result)
    except Exception:
        return False

def remove_redundant_boxes(omml_str):
    try:
        # Обернуть в корневой тег с объявлением пространства имён
        wrapped_omml = f'<wrapper xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">{omml_str}</wrapper>'
        root = ET.fromstring(wrapped_omml)
        m_ns = "{http://schemas.openxmlformats.org/officeDocument/2006/math}"

        # Упрощаем <m:oMath><m:box><m:e>...</m:e></m:box></m:oMath>
        for math in root.findall(f".//{m_ns}oMath"):
            children = list(math)
            if len(children) == 1 and children[0].tag == f"{m_ns}box":
                box = children[0]
                elems = list(box)
                if len(elems) == 1 and elems[0].tag == f"{m_ns}e":
                    e_children = list(elems[0])
                    math.clear()
                    for ec in e_children:
                        math.append(ec)

        # Рекурсивно разворачиваем все box → e → *
        def recursive_unwrap(parent):
            for i, child in enumerate(list(parent)):
                recursive_unwrap(child)
                if child.tag == f"{m_ns}box":
                    elems = list(child)
                    if len(elems) == 1 and elems[0].tag == f"{m_ns}e":
                        e = elems[0]
                        e_children = list(e)
                        parent.remove(child)
                        for j, ec in enumerate(e_children):
                            parent.insert(i + j, ec)

        recursive_unwrap(root)

        # Возврат только внутреннего содержимого (без <wrapper>)
        inner = "".join(ET.tostring(e, encoding="unicode") for e in root)
        return inner
    except Exception as e:
        print(f"[remove_redundant_boxes ERROR]: {type(e).__name__}: {e}")
        return omml_str
    
def protect_bra_ket(latex):
    # Заменяет все конструкции |...⟩ (например, |0\rangle) на уникальные токены
    pattern = r'(\|[^\|{}\\]+\s*\\rangle)'
    braket_dict = {}
    def replacer(match):
        key = f"__BRAKET_{len(braket_dict)}__"
        braket_dict[key] = match.group(1)
        return key
    latex = re.sub(pattern, replacer, latex)
    return latex, braket_dict

def restore_bra_ket(latex, braket_dict):
    for key, value in braket_dict.items():
        latex = latex.replace(key, value)
    return latex
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_horizontal_line(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)


def fix_integral_indices(latex: str) -> str:
    latex = re.sub(r'(\\(?:int|oint|iint|iiint|idotsint))\\limits', r'\1', latex)

    # 1. Есть только нижний индекс (фигурные скобки)
    latex = re.sub(
        r'(\\(?:int|oint|iint|iiint|idotsint))\s*_\{([^{}]+)\}(?!\s*\^)',
        lambda m: f"{m.group(1)}_{{{m.group(2)}}}^{{\\ }}",
        latex
    )

    # 2. Есть только нижний индекс (без фигурных скобок)
    latex = re.sub(
        r'(\\(?:int|oint|iint|iiint|idotsint))_([A-Za-zA-Z0-9\\]+)(?!\s*\^)',
        lambda m: f"{m.group(1)}_{{{m.group(2)}}}^{{\\ }}",
        latex
    )

    # 3. Есть только верхний индекс (фигурные скобки)
    latex = re.sub(
        r'(\\(?:int|oint|iint|iiint|idotsint))\s*\^\{([^{}]+)\}(?!\s*_|\{)',
        lambda m: f"{m.group(1)}^{{{m.group(2)}}}_{{\\ }}",
        latex
    )

    # 4. Есть только верхний индекс (без фигурных скобок)
    latex = re.sub(
        r'(\\(?:int|oint|iint|iiint|idotsint))\^([A-Za-zA-Z0-9\\]+)(?!\s*_|\{)',
        lambda m: f"{m.group(1)}^{{{m.group(2)}}}_{{\\ }}",
        latex
    )

    # 5. Нет ни одного индекса — добавить оба
    latex = re.sub(
        r'(\\(?:int|oint|iint|iiint|idotsint))(?![\s_^{])',
        lambda m: f"{m.group(1)}_{{\\ }}^{{\\ }}",
        latex
    )

    return latex

def fix_sum_indices(latex: str) -> str:
    # Ищет \sum, после которого есть нижний индекс (_{...} или _a), но нет ^{...} или ^a сразу после
    pattern = r'(\\sum\s*_(?:\{[^{}]*\}|[^\s\^{}]+))(?!\s*\^)'
    def repl(match):
        return match.group(1) + '^{\\infty}'
    return re.sub(pattern, repl, latex)

def find_text_block_ends(latex):
    pattern = r'\\text[a-zA-Z]*\{((?:[^{}]+|(?R))*)\}'
    return [m.end() for m in regex.finditer(pattern, latex)]

def wrap_func_args(latex):
    funcs = ['sin', 'cos', 'tan', 'cot', 'sec', 'csc',
             'arcsin', 'arccos', 'arctan',
             'sinh', 'cosh', 'tanh',
             'log', 'ln', 'exp']
    
    # Поддержка форм: \func x  → \func(x), но НЕ трогать если уже \func{...} или \func\left(...)
    pattern = r'(\\(?:' + '|'.join(funcs) + r')(\^\d+)?)(\s*)([a-zA-Z0-9\\]+)'

    def repl(m):
        func = m.group(1)
        arg = m.group(4)

        # Пропускаем, если аргумент уже обёрнут (например, \left..., {..., \frac...)
        if arg.startswith('\\left') or arg.startswith('\\frac') or arg.startswith('{') or arg.startswith('\\begin'):
            return m.group(0)
        
        return f"{func}({arg})"

    prev = None
    while prev != latex:
        prev = latex
        latex = re.sub(pattern, repl, latex)

    return latex

def add_space_number_command(formula):
    # Между числом и командой (2\pi), между числом и буквой (2x)
    formula = re.sub(r'(\d)(\\[a-zA-Z]+)', r'\1 \2', formula)
    formula = re.sub(r'(\d)([a-zA-Z])', r'\1 \2', formula)
    return formula

def insert_multiplication_dots_smart(latex: str, enable_log=False) -> str:

    LATEX_GREEK = [
        'alpha','beta','gamma','delta','epsilon','zeta','eta','theta','vartheta',
        'iota','kappa','lambda','mu','nu','xi','omicron','pi','rho','sigma',
        'tau','upsilon','phi','varphi','chi','psi','omega','varepsilon','varsigma',
        'varrho'
    ]
    LATEX_FUNCTIONS = [
        'sin', 'cos', 'tan', 'cot', 'sec', 'csc',
        'arcsin', 'arccos', 'arctan',
        'sinh', 'cosh', 'tanh',
        'log', 'ln', 'exp'
    ]
    FORBIDDEN_RELATIONS = ['\\leq', '\\geq', '\\neq', '\\approx', '=', '<', '>', '\\to']
    FORBIDDEN_QUANTORS = ['\\forall', '\\exists']
    FORBIDDEN_SET_OPERATORS = ['\\in', '\\notin', '\\subset', '\\supset', '\\subseteq', '\\supseteq']
    FORBIDDEN_SETS = ['\\mathbb{R}', '\\mathbb{Z}', '\\mathbb{N}', '\\mathbb{Q}', '\\mathbb{C}']
    MATH_OPS = [
    '\\int', '\\oint', '\\iint', '\\iiint', '\\idotsint',
    '\\sum', '\\prod', '\\lim', '\\limsup', '\\liminf', '\\bigsqcup',
    '\\bigcap', '\\bigcup', '\\bigvee', '\\bigwedge', '\\bigodot',
    '\\bigotimes', '\\bigoplus', '\\biguplus'
    ]

    def is_greek(token):
        return re.fullmatch(r'\\(' + '|'.join(LATEX_GREEK) + r')', token)
    def is_function(token):
        m = re.fullmatch(r'\\([a-zA-Z]+)', token)
        return m and m.group(1) in LATEX_FUNCTIONS
    def is_latin(token):
        return re.fullmatch(r'[a-zA-Z]', token)
    def is_number(token):
        return re.fullmatch(r'[0-9]+([.,][0-9]+)?', token)
    def is_special(token):
        return token.startswith('__') and token.endswith('__')
    def is_relation(token):
        return token in FORBIDDEN_RELATIONS
    def is_quantor(token):
        return token in FORBIDDEN_QUANTORS
    def is_set_op(token):
        return token in FORBIDDEN_SET_OPERATORS
    def is_set(token):
        return token in FORBIDDEN_SETS

    # 1. Преобразовать \quad, \qquad и любые последовательности пробелов в один спец-символ-разделитель:
    latex = re.sub(r'(\\quad|\\qquad|\s+)', ' ␣ ', latex)  # символ "␣" как технический разделитель
    tokens = [t for t in re.split(r'(\s+|␣)', latex) if t and not t.isspace()]

    result = []
    i = 0
    while i < len(tokens):
        token = tokens[i]

        if token == '␣' and (i + 1 >= len(tokens) or tokens[i + 1] == '␣'):
            i += 1
            continue
        if token == '␣':
            # Проверим предыдущий и следующий токен на whitelist
            if result and i + 1 < len(tokens):
                lhs = result[-1]
                rhs = tokens[i + 1]
                # Не вставлять ничего если один из токенов спецблок/знак/квантор/отношение
                if any([
                    is_special(lhs), is_special(rhs),
                    lhs in ['(', '[', '{', '|'] or rhs in [')', ']', '}', '|'],
                    is_relation(lhs), is_relation(rhs),
                    is_quantor(lhs), is_quantor(rhs),
                    is_set_op(lhs), is_set_op(rhs),
                    is_set(lhs), is_set(rhs),
                    (lhs in ['\\partial', 'd'] and is_latin(rhs))
                ]):
                    pass  # Просто ничего не вставлять (ни пробел, ни умножение)
                elif (
                    (is_latin(lhs) or is_greek(lhs) or is_number(lhs) or re.fullmatch(r'\\([a-zA-Z]+)', lhs))
                    and rhs in MATH_OPS
                ):
                    result.append('\\cdot')
                # Вставлять умножение по whitelist
                elif (
                    is_latin(lhs) and is_latin(rhs) or
                    is_number(lhs) and is_latin(rhs) or
                    is_latin(lhs) and is_number(rhs) or
                    is_number(lhs) and is_greek(rhs) or
                    (is_latin(lhs) and is_greek(rhs)) or (is_greek(lhs) and is_latin(rhs)) or
                    is_greek(lhs) and is_greek(rhs) or
                    (re.fullmatch(r'\\([a-zA-Z]+)', lhs) and not is_function(lhs) and is_latin(rhs) and lhs not in FORBIDDEN_RELATIONS)
                ):
                    if result and result[-1] == '\\cdot':
                        pass
                    else:
                        result.append('\\cdot')

                # Остальное — просто ничего
            i += 1  # пропускаем '␣'
        else:
            result.append(token)
            i += 1

    latex = ' '.join(result)
    return latex.strip()


def extract_matrices(latex):
    matrices = {}
    # 1. Ловим "голые" матрицы, как раньше
    matrix_pattern = r'(\\begin\{(?:pmatrix|bmatrix|matrix|smallmatrix|array)\}.*?\\end\{(?:pmatrix|bmatrix|matrix|smallmatrix|array)\})'
    # 2. Ловим матрицу ВМЕСТЕ с внешними \left...\right...
    #   Например: \left| ... \begin{...} ... \end{...} ... \right|
    #   или      \left( ... \begin{...} ... \end{...} ... \right)
    #   (опционально допускаем пробелы между \left... и \begin)
    bracket_matrix_pattern = (
        r'(\\left[\(\[\|\.]\s*'
        r'(?:\\begin\{(?:pmatrix|bmatrix|matrix|smallmatrix|array)\}.*?\\end\{(?:pmatrix|bmatrix|matrix|smallmatrix|array)\})'
        r'\s*\\right[\)\]\|\.])'
    )
    # Сначала ловим с внешними скобками (иначе "голый" блок сработает раньше)
    def bracket_matrix_replacer(match):
        key = f"__MATRIX_BLOCK_{len(matrices)}__"
        matrices[key] = match.group(0)
        return key

    latex = re.sub(bracket_matrix_pattern, bracket_matrix_replacer, latex, flags=re.DOTALL)
    # Затем "голые" блоки
    def matrix_replacer(match):
        key = f"__MATRIX_BLOCK_{len(matrices)}__"
        matrices[key] = match.group(0)
        return key

    latex = re.sub(matrix_pattern, matrix_replacer, latex, flags=re.DOTALL)
    return latex, matrices


def restore_matrices(latex, matrices):
    for key, block in matrices.items():
        latex = latex.replace(key, block)
    return latex

def replace_sqrt_with_explicit_degree(latex: str) -> str:
    # Поддержка вложенных скобок
    pattern = r'\\sqrt\{((?:[^{}]+|(?R))*)\}'
    repl = r'\\sqrt[2]{\1}'
    result = regex.sub(pattern, repl, latex)
    return regex.sub(pattern, repl, latex)

def add_sqrt_degree(latex):
    result = ''
    i = 0
    while i < len(latex):
        if latex[i:i+5] == '\\sqrt':
            j = i + 5
            # Пропускаем пробелы после \sqrt
            while j < len(latex) and latex[j] == ' ':
                j += 1
            # Если следующая скобка — это [, то степень уже есть
            if j < len(latex) and latex[j] == '[':
                result += latex[i:j+1]  # копируем до [
                i = j + 1
                continue
            # Если сразу после \sqrt идет {
            elif j < len(latex) and latex[j] == '{':
                result += '\\sqrt[2]{'
                i = j + 1
                continue
            else:
                # просто копируем \sqrt и движемся дальше
                result += latex[i]
                i += 1
                continue
        else:
            result += latex[i]
            i += 1
    return result

def pad_text_inside_text_blocks(latex: str) -> str:
    return re.sub(
        r'\\text\{([^}]+)\}',
        lambda m: f'\\text{{ {m.group(1).strip()} }}',
        latex
    )

def clean_latex(formula: str) -> str:
    f = formula.strip()
    while f.startswith('$'):
        f = f[1:]
    while f.endswith('$'):
        f = f[:-1]
    f = f.strip()
    if f.endswith('\\'):
        f = f[:-1].rstrip()
    return f



def auto_bracket_dot_power(latex):
    result = []
    i = 0
    n = len(latex)
    while i < n:
        # ищем \dot{...}^ или \ddot{...}^
        if latex.startswith(r'\dot{', i) or latex.startswith(r'\ddot{', i):
            is_ddot = latex.startswith(r'\ddot{', i)
            tag = r'\ddot{' if is_ddot else r'\dot{'
            tag_len = len(tag)
            result.append('(')  # открывающая скобка
            result.append(tag)
            i += tag_len
            k = 1
            while i < n:
                c = latex[i]
                result.append(c)
                if c == '{':
                    k += 1
                elif c == '}':
                    k -= 1
                    if k == 0:
                        i += 1
                        break
                i += 1
            # после закрывающей }
            # возможны пробелы
            while i < n and latex[i].isspace():
                result.append(latex[i])
                i += 1
            # если дальше идёт ^
            if i < n and latex[i] == '^':
                result.append(')')
                result.append('^')
                i += 1
            else:
                # не степень — просто закрываем скобку
                result.append(')')
        else:
            result.append(latex[i])
            i += 1
    return ''.join(result)

def latex_to_omml(latex: str) -> str:
    #with open("debug_latex_trace.txt", "a", encoding="utf-8") as log:
        #def logstep(title, content):
            #log.write(f"--- {title} ---\n{content}\n\n")

        # Вырезаем \tag{...} из конца формулы
        tag_text = None
        m = re.search(r'\\tag\{([^{}]+)\}\s*$', latex)
        if m:
            tag_text = m.group(1).strip()
            latex = latex[:m.start()].rstrip()

        latex = auto_bracket_dot_power(latex)
        #logstep("01 auto_bracket_dot_power", latex)

        latex_fixed = re.sub(r'\\vec\{([^}]+)\}', r'\\overset{\\rightarrow}{\1}', latex)
        #logstep("02 replace vec", latex_fixed)

        latex_fixed = add_sqrt_degree(latex_fixed)
        #logstep("03 add_sqrt_degree", latex_fixed)

        latex_fixed = wrap_func_args(latex_fixed)
        #logstep("04 wrap_func_args", latex_fixed)

        latex_fixed = fix_integral_indices(latex_fixed)
        #logstep("05 fix_integral_indices", latex_fixed)

        latex_fixed = fix_sum_indices(latex_fixed)
        #logstep("06 fix_sum_indices", latex_fixed)

        latex_fixed = add_space_number_command(latex_fixed)
        #logstep("07 add_space_number_command", latex_fixed)

        latex_fixed, braket_dict = protect_bra_ket(latex_fixed)
        #logstep("08 protect_bra_ket", latex_fixed)

        latex_fixed = insert_multiplication_dots_smart(latex_fixed)
        #logstep("09 insert_multiplication_dots_smart", latex_fixed)

        latex_fixed = restore_bra_ket(latex_fixed, braket_dict)
        #logstep("10 restore_bra_ket", latex_fixed)

        latex_fixed = pad_text_inside_text_blocks(latex_fixed)
        #logstep("11 pad_text_inside_text_blocks", latex_fixed)

        latex_fixed = re.sub(r'\\+,', r'\\THINSPACE ', latex_fixed)
        #logstep("12 before mathml convert", latex_fixed)


        # Защита \tag
        latex_fixed = re.sub(r'\\tag\{([^{}]+)\}', r'\\TAGSTART\1\\TAGEND', latex_fixed)

        mathml = latex2mathml.converter.convert(latex_fixed)
        # Восстановление \tag
        mathml = mathml.replace(r'\TAGSTART', r'\tag{').replace(r'\TAGEND', r'}')

        mathml = mathml.replace(r'<mi>\THINSPACE</mi>', r'<mspace width="1.0em"/>')
        mathml = mathml.replace(r'<mo>\THINSPACE</mo>', r'<mspace width="1.0em"/>')

        omml = mathml2omml.convert(mathml, html.entities.name2codepoint)
        omml = remove_redundant_boxes(omml)
        #logstep("14 omml", omml)

        return omml, tag_text

def parse_markdown_styles(text):
    markers = [
        ('***', ['bold', 'italic']),
        ('**', ['bold']),
        ('*', ['italic']),
        ('~~', ['strike']),
        ('__', ['bold']),
        ('_', ['italic']),
    ]
    markers.sort(key=lambda x: -len(x[0]))
    result = []
    stack = []
    i = 0
    buf = ''
    while i < len(text):
        found = False
        for marker, styles in markers:
            if text[i:i+len(marker)] == marker:
                if stack and stack[-1][0] == marker:
                    if buf:
                        stack[-1][2].append((buf, []))
                        buf = ''
                    complete = stack.pop()
                    item = (complete[2], complete[1])
                    if stack:
                        stack[-1][2].append(item)
                    else:
                        result.append(item)
                    i += len(marker)
                    found = True
                    break
                else:
                    if buf:
                        if stack:
                            stack[-1][2].append((buf, []))
                        else:
                            result.append((buf, []))
                        buf = ''
                    stack.append([marker, styles, []])
                    i += len(marker)
                    found = True
                    break
        if not found:
            buf += text[i]
            i += 1
    if buf:
        if stack:
            stack[-1][2].append((buf, []))
        else:
            result.append((buf, []))
    while stack:
        content = stack.pop()
        if content[2]:
            if stack:
                stack[-1][2].extend(content[2])
            else:
                result.extend(content[2])
    return result

def add_markdown_runs(paragraph, parsed, parent_styles=None):
    if parent_styles is None:
        parent_styles = []
    for part, styles in parsed:
        all_styles = list(parent_styles) + list(styles)
        if isinstance(part, list):
            add_markdown_runs(paragraph, part, all_styles)
        else:
            run = paragraph.add_run(part)
            if 'bold' in all_styles:
                run.bold = True
            if 'italic' in all_styles:
                run.italic = True
            if 'strike' in all_styles:
                run.font.strike = True

def add_omml_run(paragraph, omml: str):
    run_xml = (
        '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
        f"{omml}</w:r>"
    )
    run_element = parse_xml(run_xml)
    paragraph._element.append(run_element)

def process_inline_markdown(paragraph, text):
    runs = []
    last = 0
    pattern = r'(?<!\\)(\$[^$]+\$|\\\((.*?)\\\)|\\\[(.*?)\\\])'
    for m in re.finditer(pattern, text):
        if m.start() > last:
            runs.append(("text", text[last:m.start()]))
        raw = m.group(0)
        if raw.startswith('$') and raw.endswith('$'):
            formula = clean_latex(raw[1:-1])
        elif raw.startswith('\\(') and raw.endswith('\\)'):
            formula = clean_latex(raw[2:-2])
        elif raw.startswith('\\[') and raw.endswith('\\]'):
            formula = clean_latex(raw[2:-2])
        else:
            formula = clean_latex(raw)
        runs.append(("math", formula))
        last = m.end()
    if last < len(text):
        runs.append(("text", text[last:]))
    elif not runs:
        runs.append(("text", text))

    for kind, content in runs:
        if kind == "math":
            try:
                omml, tag = latex_to_omml(content)
                add_omml_run(paragraph, omml)
                if tag:
                    run = paragraph.add_run()
                    run.add_tab()
                    run.add_text(f"({tag})")
                    paragraph.paragraph_format.tab_stops.add_tab_stop(Pt(468))  # Пример: 6.5" = 468pt

            except Exception as e:
                paragraph.add_run(f"[Ошибка формулы: {content} ({type(e).__name__}: {e})]")
        else:
            parsed = parse_markdown_styles(content)
            add_markdown_runs(paragraph, parsed)


def markdown_to_docx(md_text, output_file):
    TAB_POS = 300  # ← Константа табуляции (3 дюйма от левого края)
    doc = Document()
    lines = md_text.split('\n')
    i = 0
    while i < len(lines):
        # --- обработка блока ```math ... ```
        if lines[i].strip() == '```math':
            formula_lines = []
            i += 1
            while i < len(lines) and lines[i].strip() != '```':
                formula_lines.append(lines[i])
                i += 1
            # пропустить закрывающий ```
            if i < len(lines) and lines[i].strip() == '```':
                i += 1
            formula = clean_latex('\n'.join(formula_lines))
            p_formula = doc.add_paragraph()
            try:
                omml, tag = latex_to_omml(formula)
                add_omml_run(p_formula, omml)
                if tag:
                    run = p_formula.add_run()
                    run.add_tab()
                    run.add_text(f"({tag})")
                    p_formula.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    p_formula.paragraph_format.tab_stops.add_tab_stop(Pt(TAB_POS))
            except Exception as e:
                p_formula.add_run(f"[Ошибка формулы: {formula} ({type(e).__name__}: {e})]")
            continue
        
        if lines[i].strip().startswith(r'\['):
            formula_lines = []
            # Удалить начальный \[
            line = lines[i].lstrip()
            if line.startswith(r'\['):
                line = line[2:]
            if line.rstrip().endswith(r'\]'):
                line = re.sub(r'\\\]$', '', line.rstrip())
                formula_lines.append(line)
            else:
                formula_lines.append(line)
                i += 1
                while i < len(lines):
                    line = lines[i].rstrip()
                    if line.endswith(r'\]'):
                        formula_lines.append(line[:-2])
                        break
                    formula_lines.append(line)
                    i += 1
            formula = clean_latex('\n'.join(formula_lines))
            p_formula = doc.add_paragraph()
            try:
                omml, tag = latex_to_omml(formula)
                add_omml_run(p_formula, omml)
                if tag:
                    run = p_formula.add_run()
                    run.add_tab()
                    run.add_text(f"({tag})")
                    p_formula.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    p_formula.paragraph_format.tab_stops.add_tab_stop(Pt(TAB_POS))
            except Exception as e:
                p_formula.add_run(f"[Ошибка формулы: {formula} ({type(e).__name__}: {e})]")
            i += 1
            continue


        line = lines[i].rstrip()
        if line.strip() == "---":
            p = doc.add_paragraph()
            add_horizontal_line(p)
            i += 1
            continue
        if line.strip().startswith('$$'):
            if line.strip().endswith('$$') and len(line.strip()) > 4:
                formula = clean_latex(line)
                p_formula = doc.add_paragraph()
                try:
                    omml, tag = latex_to_omml(formula)
                    add_omml_run(p_formula, omml)
                    if tag:
                        run = p_formula.add_run()
                        run.add_tab()
                        run.add_text(f"({tag})")
                        p_formula.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                        p_formula.paragraph_format.tab_stops.add_tab_stop(Pt(TAB_POS))
                except Exception as e:
                    p_formula.add_run(f"[Ошибка формулы: {formula} ({type(e).__name__}: {e})]")
                i += 1
                continue
            formula_lines = [line]
            i += 1
            while i < len(lines) and not lines[i].strip().endswith('$$'):
                formula_lines.append(lines[i])
                i += 1
            if i < len(lines):
                formula_lines.append(lines[i])
                i += 1
            formula_block = '\n'.join(formula_lines)
            formula = clean_latex(formula_block)
            p_formula = doc.add_paragraph()
            try:
                omml, tag = latex_to_omml(formula)
                add_omml_run(p_formula, omml)
                if tag:
                    run = p_formula.add_run()
                    run.add_tab()
                    run.add_text(f"({tag})")
                    p_formula.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    p_formula.paragraph_format.tab_stops.add_tab_stop(Pt(TAB_POS))
            except Exception as e:
                p_formula.add_run(f"[Ошибка формулы: {formula} ({type(e).__name__}: {e})]")
            continue

        if not line.strip():
            doc.add_paragraph()
            i += 1
            continue

        m_header = re.match(r'^(#+)\s*(.*)', line)
        if m_header:
            level = len(m_header.group(1))
            text = m_header.group(2)
            p = doc.add_paragraph()
            process_inline_markdown(p, text)
            size = Pt(19) if level == 1 else Pt(16) if level == 2 else Pt(14) if level == 3 else Pt(13) if level == 4 else Pt(12) if level == 5 else Pt(11)
            for run in p.runs:
                run.bold = True
                run.font.size = size
            i += 1
            continue

        m_list = re.match(r'^[-*]\s+(.*)', line)
        if m_list:
            text = m_list.group(1)
            p = doc.add_paragraph(style='List Bullet')
            process_inline_markdown(p, text)
            i += 1
            continue

        # Обработка Markdown-таблицы
        if i + 1 < len(lines) and '|' in lines[i] and '|' in lines[i + 1] and set(lines[i + 1].replace('|', '').strip()) <= {'-', ' '}:
            table_lines = []
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i])
                i += 1
            cleaned_lines = [line.strip().strip('|') for line in table_lines]
            if len(cleaned_lines) >= 2 and set(cleaned_lines[1].replace('|', '').strip()) <= {'-', ' '}:
                del cleaned_lines[1]  # удаляем строку-разделитель
            csv_text = '\n'.join(cleaned_lines)
            reader = csv.reader(StringIO(csv_text), delimiter='|')
            rows = list(reader)
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = 'Table Grid'
                table.autofit = True
                table.first_row = True
                for r, row in enumerate(rows):
                    for c, cell in enumerate(row):
                        cell_text = cell.strip()
                        cell = table.cell(r, c)
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        paragraph = cell.paragraphs[0]
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        process_inline_markdown(paragraph, cell_text)
                        if r == 0 and paragraph.runs:
                            paragraph.runs[0].bold = True
            continue
        p = doc.add_paragraph()
        process_inline_markdown(p, line)
        i += 1

    doc.save(output_file)
    return True

def clean_spaces_before_dollars_and_begin_end(text):
    result = []
    n = len(text)
    i = 0
    while i < n:
        # Обработка перевода строки + пробелы перед $$, \begin, \end
        if text[i:i+2] == '\r\n':
            result.append('\r\n')
            i += 2
            start = i
            while i < n and text[i] in ' \t':
                i += 1
            if text.startswith('$$', i):
                # --- $$ блок
                result.append('\n')
                result.append('$$')
                i += 2
                if i < n and text[i] != '\n':
                    result.append('\n')
            elif text.startswith(r'\begin', i):
                result.append(r'\begin')
                i += len(r'\begin')
            elif text.startswith(r'\end', i):
                result.append(r'\end')
                i += len(r'\end')
            else:
                result.extend(text[start:i])
        elif text[i] == '\n':
            result.append('\n')
            i += 1
            start = i
            while i < n and text[i] in ' \t':
                i += 1
            if text.startswith('$$', i):
                result.append('\n')
                result.append('$$')
                i += 2
                if i < n and text[i] != '\n':
                    result.append('\n')
            elif text.startswith(r'\begin', i):
                result.append(r'\begin')
                i += len(r'\begin')
            elif text.startswith(r'\end', i):
                result.append(r'\end')
                i += len(r'\end')
            else:
                result.extend(text[start:i])
        elif text[i] == '\r':
            result.append('\r')
            i += 1
            start = i
            while i < n and text[i] in ' \t':
                i += 1
            if text.startswith('$$', i):
                result.append('\n')
                result.append('$$')
                i += 2
                if i < n and text[i] != '\n':
                    result.append('\n')
            elif text.startswith(r'\begin', i):
                result.append(r'\begin')
                i += len(r'\begin')
            elif text.startswith(r'\end', i):
                result.append(r'\end')
                i += len(r'\end')
            else:
                result.extend(text[start:i])
        # Прямо $$ (не на новой строке)
        elif text[i:i+2] == '$$':
            # Убираем пробелы до $$ (если были)
            j = len(result) - 1
            while j >= 0 and result[j] in ' \t':
                result.pop()
                j -= 1
            if j < 0 or result[j] != '\n':
                result.append('\n')
            result.append('$$')
            i += 2
            # После открывающего $$ — если нет \n, ставим
            if i < n and text[i] != '\n':
                result.append('\n')
        # После закрывающего $$
        elif len(result) >= 2 and ''.join(result[-2:]) == '$$':
            # Ставим \n если его нет (но не дублируем)
            if text[i] not in '\n\r':
                result.append('\n')
            result.append(text[i])
            i += 1
        else:
            result.append(text[i])
            i += 1
    return ''.join(result)


def replace_newlines_manual(text):
    def startswith_any(text, i, patterns):
        for p in patterns:
            if text.startswith(p, i):
                return True
        return False

    matrix_begins = [
        r'\begin{pmatrix}',
        r'\begin{bmatrix}',
        r'\begin{matrix}',
        r'\begin{smallmatrix}',
        r'\begin{vmatrix}',
        r'\begin{Vmatrix}',
        r'\begin{array}',
    ]

    result = []
    i = 0
    text = clean_spaces_before_dollars_and_begin_end(text)
    n = len(text)
    while i < n:
        if text[i:i+2] == '$$':
            result.append('$$')
            i += 2

            # Пропускаем все пробелы, табы и переводы строк после $$
            while i < n and text[i] in ' \t\r\n':
                i += 1

            # Если следующий блок - \begin{...}, то пропускаем весь блок до следующего $$
            if startswith_any(text, i, matrix_begins):
                block_start = i
                next_dollars = text.find('$$', block_start)
                if next_dollars != -1:
                    result.append(text[block_start:next_dollars])
                    i = next_dollars
                    result.append('$$')
                    i += 2
                    continue

            # Собираем формулу в буфер
            buf = []
            formula_has_matrix = False
            preview_j = i
            # Проверяем наличие матричных конструкций в этом блоке (до следующего $$)
            next_dollars = text.find('$$', i)
            if next_dollars == -1:
                next_dollars = n
            preview_block = text[i:next_dollars]
            for m in matrix_begins:
                if m in preview_block:
                    formula_has_matrix = True
                    break

            early_break = False
            while i < n:
                # Конец формулы (с учетом разных переносов)
                if (text[i:i+2] == '\r\n' and text[i+2:i+4] == '$$'):
                    result.append(''.join(buf))
                    result.append('\r\n')
                    i += 2
                    early_break = True
                    break
                if (text[i] == '\n' and text[i+1:i+3] == '$$'):
                    result.append(''.join(buf))
                    result.append('\n')
                    i += 1
                    early_break = True
                    break
                if (text[i] == '\r' and text[i+1:i+3] == '$$'):
                    result.append(''.join(buf))
                    result.append('\r')
                    i += 1
                    early_break = True
                    break
                if text[i:i+2] == '$$':
                    break
                if text[i] == '\\' and i + 1 < n and text[i + 1] == '$':
                    buf.append('\\$')
                    i += 2
                    continue
                if text[i:i+2] == '\r\n':
                    if formula_has_matrix:
                        buf.append(' \\ ')
                    else:
                        buf.append(' ')
                    i += 2
                elif text[i] == '\n' or text[i] == '\r':
                    if formula_has_matrix:
                        buf.append(' \\ ')
                    else:
                        buf.append(' ')
                    i += 1
                else:
                    buf.append(text[i])
                    i += 1
            if not early_break:
                result.append(''.join(buf))
            if text[i:i+2] == '$$':
                result.append('$$')
                i += 2
        elif text[i] == '$':
            # Аналогичная логика для inline формул, если надо
            result.append('$')
            i += 1
            buf = []
            # Проверяем наличие матриц внутри $...$
            formula_has_matrix = False
            preview_j = i
            next_dollar = text.find('$', i)
            if next_dollar == -1:
                next_dollar = n
            preview_block = text[i:next_dollar]
            for m in matrix_begins:
                if m in preview_block:
                    formula_has_matrix = True
                    break
            while i < n and text[i] != '$':
                if text[i] == '\\' and i + 1 < n and text[i + 1] == '$':
                    buf.append('\\$')
                    i += 2
                    continue
                if text[i:i+2] == '\r\n':
                    if formula_has_matrix:
                        buf.append(' \\ ')
                    else:
                        buf.append(' ')
                    i += 2
                elif text[i] == '\n' or text[i] == '\r':
                    if formula_has_matrix:
                        buf.append(' \\ ')
                    else:
                        buf.append(' ')
                    i += 1
                else:
                    buf.append(text[i])
                    i += 1
            result.append(''.join(buf))
            if i < n and text[i] == '$':
                result.append('$')
                i += 1
        else:
            result.append(text[i])
            i += 1

    return ''.join(result)


def fix_inline_formula_spaces(text):
    # Добавить пробелы вокруг одиночных inline $...$
    # Пробел перед $, если нет пробела или начала строки
    text = re.sub(r'(?<![\s$])\$([^\$].*?[^\$])\$(?![\s$])', r' $\1$ ', text)
    return text


def main():
    if len(sys.argv) >= 2:
        input_file = sys.argv[1]
        if len(sys.argv) >= 3:
            output_file = sys.argv[2]
        else:
            output_file = "output.docx"
        try:
            with open(input_file, 'r', encoding='utf-8') as f:
                md_text = f.read()
            md_text = fix_inline_formula_spaces(md_text)
            md_text = replace_newlines_manual(md_text)
            ok = markdown_to_docx(md_text, output_file)
            if ok:
                print(f"Conversion complete.\nFile saved as: {output_file}")
            else:
                print("File is not saved.")
        except Exception as e:
            print(f"[Error] Could not open or read the file: '{input_file}': {e}")
            sys.exit(1)
    else:
        print(f"aimath2docx.py input.txt output.docx")

if __name__ == "__main__":
    main()
